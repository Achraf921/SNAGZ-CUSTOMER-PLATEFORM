from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json
import sys
import os
import base64
import datetime
import tempfile
import zipfile
import re
import shutil

# Set up logging to a file
def setup_logging():
    log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
    os.makedirs(log_dir, exist_ok=True)
    log_file = os.path.join(log_dir, f'docx_processor_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.log')
    
    def log_message(message):
        with open(log_file, 'a', encoding='utf-8') as f:
            timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            f.write(f'[{timestamp}] {message}\n')
            print(f'[{timestamp}] {message}')
    
    return log_message

log = setup_logging()


def replace_placeholders_and_format(docx_path, shop_data, output_path):
    log(f"Starting document processing for: {docx_path}")
    log(f"Output will be saved to: {output_path}")
    
    try:
        document = Document(docx_path)
        log(f"Successfully loaded document: {docx_path}")
    except Exception as e:
        log(f"Failed to load document: {str(e)}")
        raise

    log(f"[DEBUG - Inside replace_placeholders_and_format] Type of shop_data: {type(shop_data)}")
    log(f"[DEBUG - Inside replace_placeholders_and_format] Content of shop_data: {shop_data}")
    
    # Ensure shop_data is a dictionary
    if not isinstance(shop_data, dict):
        error_msg = f"[ERROR] shop_data is not a dictionary. Type: {type(shop_data)}, Value: {shop_data}"
        log(error_msg)
        raise ValueError(error_msg)

    # Define the mapping from XXXn to shop_data keys
    # Values are now pre-formatted from the backend
    
    # Function to clean Unicode zero-width characters that might interfere with replacement
    def clean_text_for_replacement(text):
        # Remove common zero-width Unicode characters
        text = text.replace('\u200b', '')  # Zero-width space
        text = text.replace('\u200c', '')  # Zero-width non-joiner  
        text = text.replace('\u200d', '')  # Zero-width joiner
        text = text.replace('\u2060', '')  # Word joiner
        text = text.replace('\ufeff', '')  # Zero-width no-break space
        text = text.replace('\u202f', '')  # Narrow no-break space
        text = text.replace('\u00a0', '')  # Non-breaking space
        text = text.replace('\u2009', '')  # Thin space
        text = text.replace('\u200a', '')  # Hair space
        text = text.replace('\u2028', '')  # Line separator
        text = text.replace('\u2029', '')  # Paragraph separator
        return text
    
    # More aggressive replacement function that handles Unicode around placeholders
    def replace_placeholder_with_unicode_handling(text, placeholder, value):
        # First try normal replacement
        if placeholder in text:
            return text.replace(placeholder, str(value))
        
        # If normal replacement didn't work, try with Unicode cleaning
        cleaned = clean_text_for_replacement(text)
        if placeholder in cleaned:
            # Find all possible Unicode variations around the placeholder
            import re
            # Create a pattern that matches the placeholder with any Unicode characters around it
            pattern = r'[\u200b\u200c\u200d\u2060\ufeff\u202f\u00a0\u2009\u200a\u2028\u2029]*' + re.escape(placeholder) + r'[\u200b\u200c\u200d\u2060\ufeff\u202f\u00a0\u2009\u200a\u2028\u2029]*'
            return re.sub(pattern, str(value), text)
        
        return text
    
    placeholder_mapping = {
        "XXX1": shop_data.get("nomProjet", ""),  # XXX1: nomProjet
        "XXX2": shop_data.get("typeProjet", ""),  # XXX2: typeProjet
        "XXX3": shop_data.get("commercial", ""),  # XXX3: commercial (salesperson)
        "XXX4": shop_data.get("nomClient", ""),  # XXX4: nomClient (customer raisonSociale)
        "XXX5": shop_data.get("compteClientRef", ""),  # XXX5: compteClientRef (customer CompteClientNumber)
        "XXX7": shop_data.get("dateMiseEnLigne", ""),  # XXX7: dateMiseEnLigne (if available)
        "XXX8": shop_data.get("dateCommercialisation", ""),  # XXX8: dateCommercialisation (YYYY/MM/DD)
        "XXX9": shop_data.get("dateSortieOfficielle", ""),  # XXX9: dateSortieOfficielle (YYYY/MM/DD)
        "XXX10": shop_data.get("precommande", ""),  # XXX10: precommande (OUI/NON)
        "XXX11": shop_data.get("dedicaceEnvisagee", ""),  # XXX11: dedicaceEnvisagee (OUI/NON)
        "XXX12": shop_data.get("estBoutiqueEnLigne", ""),  # XXX12: estBoutiqueEnLigne (OUI/NON)
        "XXX13": shop_data.get("chefProjet", ""),  # XXX13: prenomChefProjet + ' ' + nomChefProjet
        "XXX14": shop_data.get("demarrageProjet", ""),  # XXX14: demarrageProjet (YYYY/MM/DD)
        "XXX15": shop_data.get("contactsClient", ""),  # XXX15: contactsClient (client email)
        "XXX69": shop_data.get("pourcentageSNA", ""),  # XXX69: pourcentageSNA value
        "COMPTENUM": shop_data.get("compteClientRef", ""),  # COMPTENUM: customer CompteClientNumber
    }

    # -------------------------------------------------------------
    # Dynamic strikethrough rules for Contract-D2C cost lines
    # -------------------------------------------------------------
    type_ab = str(shop_data.get("typeAbonnementShopify", "")).strip().lower()
    strike_mensuel = type_ab in ("", "aucun", "annuel")  # cross out mensuel when not chosen or when annual chosen
    strike_annuel  = type_ab in ("", "aucun", "mensuel") # cross out annuel when not chosen or when monthly chosen

    strike_mondial_relay = not bool(shop_data.get("moduleMondialRelay", False))
    strike_delivengo     = not bool(shop_data.get("moduleDelivengo", False))

    contract_d2c_strikethrough_rules = [
        {"text": "Abonnement SHOPIFY mensuel sans engagement = 88€", "strike": strike_mensuel},
        {"text": "Abonnement SHOPIFY 12 mois = 948€", "strike": strike_annuel},
        {"text": "Les coûts pour ajouter le module Mondial Relay = 34€", "strike": strike_mondial_relay},
        {"text": "Les coûts pour ajouter le module Delivengo = 34€", "strike": strike_delivengo},
    ]

    # First pass: Replace all placeholders in all runs globally
    log("Starting global placeholder replacement...")
    all_runs = []
    
    # Collect all runs from paragraphs
    for paragraph in document.paragraphs:
        all_runs.extend(paragraph.runs)
    
    # Collect all runs from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_runs.extend(paragraph.runs)
    
    # Replace placeholders in all runs
    total_runs = len(all_runs)
    log(f"Total runs to process: {total_runs}")
    
    # Log all runs with text for debugging
    text_runs = [run for run in all_runs if run.text.strip()]
    log(f"Runs with non-empty text: {len(text_runs)}")
    for i, run in enumerate(text_runs[:10]):  # Log first 10 runs
        log(f"Run {i}: '{run.text}'")
    
    for key, value in placeholder_mapping.items():
        replacements_made = 0
        log(f"Looking for placeholder: {key} (value: '{value}')")
        
        for run in all_runs:
            if run.text:  # Only process runs with text
                cleaned_run = clean_text_for_replacement(run.text).strip()
                if cleaned_run == key:  # Replace only when the entire run equals the placeholder
                    original_text = run.text
                    run.text = str(value)
                    replacements_made += 1
                    log(f"✅ (run exact) Replaced {key} with '{value}': '{original_text}' → '{run.text}'")
        
        log(f"Total replacements made for {key}: {replacements_made}")
    
    # Additional method: try replacing across paragraph text directly
    log("Attempting direct paragraph text replacement...")
    for paragraph in document.paragraphs:
        if paragraph.text:
            original_para_text = paragraph.text
            for key, value in placeholder_mapping.items():
                if key in original_para_text or key in clean_text_for_replacement(original_para_text):
                    log(f"Found {key} in paragraph: '{original_para_text}'")
                    # Try to replace by reconstructing the paragraph
                    new_para_text = replace_placeholder_with_unicode_handling(original_para_text, key, str(value))
                    if new_para_text != original_para_text:
                        # Clear existing runs and create new one
                        for run in paragraph.runs:
                            run.clear()
                        paragraph.runs[0].text = new_para_text
                        log(f"✅ Replaced in paragraph: '{original_para_text}' → '{new_para_text}'")
    
    # Additional method: try replacing in table cells directly
    log("Attempting direct table cell text replacement...")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        original_para_text = paragraph.text
                        for key, value in placeholder_mapping.items():
                            if key in original_para_text or key in clean_text_for_replacement(original_para_text):
                                log(f"Found {key} in table cell: '{original_para_text}'")
                                # Try to replace by reconstructing the paragraph
                                new_para_text = replace_placeholder_with_unicode_handling(original_para_text, key, str(value))
                                if new_para_text != original_para_text:
                                    # Clear existing runs and create new one
                                    for run in paragraph.runs:
                                        run.clear()
                                    if paragraph.runs:
                                        paragraph.runs[0].text = new_para_text
                                    else:
                                        paragraph.add_run(new_para_text)
                                    log(f"✅ Replaced in table cell: '{original_para_text}' → '{new_para_text}'")
        
        # Strikethrough application
        for rule in contract_d2c_strikethrough_rules:
            contract_text = rule["text"]
            strike = rule["strike"]

            if strike and contract_text in paragraph.text:
                for run in paragraph.runs:
                    if contract_text in run.text:
                        run.font.strike = True
    
    # Second pass: Apply strikethrough to contract D2C items in tables (if applicable)
    log("Starting strikethrough application...")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    
                    # Strikethrough application in table cells
                    for rule in contract_d2c_strikethrough_rules:
                        contract_text = rule["text"]
                        strike = rule["strike"]

                        if strike and contract_text in paragraph.text:
                            for run in paragraph.runs:
                                if contract_text in run.text:
                                    run.font.strike = True

    document.save(output_path)
    print(f"Document saved to {output_path}")

    # Perform XML-level replacement for any placeholders still not caught
    xml_level_replace(output_path, placeholder_mapping)

    log("Document processing completed successfully")

def build_cross_tag_pattern(placeholder: str) -> str:
    """Return a regex that matches the placeholder even if arbitrary XML tags or zero-width characters are interleaved
    between its characters (e.g. <w:t>XX</w:t><w:t>X</w:t><w:t>1</w:t>)."""
    # Characters that may appear between pieces of the placeholder (XML tags or invisible chars)
    gap = r"(?:<[^>]+>|\s|\u200b|\u200c|\u200d|\u2060|\ufeff|\u202f|\u00a0|\u2009|\u200a|\u2028|\u2029)*"
    # Join every character in the placeholder with the gap pattern
    escaped_chars = [re.escape(ch) for ch in placeholder]
    return gap.join(escaped_chars)


def xml_level_replace(docx_path: str, mapping: dict):
    """Open an existing DOCX file and replace placeholders at the raw XML level using robust regexes."""
    try:
        with zipfile.ZipFile(docx_path, 'r') as zin:
            # Copy all original entries but potentially modified XML files into a new zip
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_fp:
                temp_path = temp_fp.name
            with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename.endswith('.xml'):
                        xml_str = data.decode('utf-8', errors='ignore')
                        original_xml = xml_str
                        # Replace longer placeholders first to avoid partial matches (e.g. XXX1 inside XXX12)
                        for key in sorted(mapping.keys(), key=len, reverse=True):
                            value = mapping[key]
                            pattern = build_cross_tag_pattern(key)
                            xml_str, subs = re.subn(pattern, str(value), xml_str)
                            if subs:
                                log(f"[XML] Replaced {key} -> '{value}' ({subs} occurrence(s)) in {item.filename}")
                        data = xml_str.encode('utf-8')
                    zout.writestr(item, data)
        # Replace original file with the updated one
        shutil.move(temp_path, docx_path)
        log("XML-level placeholder replacement completed.")
    except Exception as e:
        log(f"[XML] Error during XML-level replacement: {e}")

if __name__ == "__main__":
    try:
        log("Script started")
        
        if len(sys.argv) != 4:
            error_msg = "Usage: python docx_processor.py <docx_template_path> <shop_data_json_string> <output_docx_path>"
            log(error_msg)
            print(error_msg, file=sys.stderr)
            sys.exit(1)

        template_path = sys.argv[1]
        encoded_shop_data_string = sys.argv[2]
        output_path = sys.argv[3]

        log(f"Template path: {template_path}")
        log(f"Output path: {output_path}")
        log(f"Encoded data length: {len(encoded_shop_data_string)} characters")

        try:
            # Decode the Base64 string back to JSON
            log("Decoding base64 data...")
            shop_data_string = base64.b64decode(encoded_shop_data_string).decode('utf-8')
            log("Successfully decoded base64 data")
            
            log(f"[DEBUG] Type of shop_data_string after decode: {type(shop_data_string)}")
            log(f"[DEBUG] First 100 chars of shop_data_string: {shop_data_string[:100]}...")

            # Parse the JSON string
            log("Parsing JSON data...")
            shop_data = json.loads(shop_data_string)
            log("Successfully parsed JSON data")
            
            log(f"[DEBUG] Type of shop_data after json.loads: {type(shop_data)}")
            log(f"[DEBUG] shop_data keys: {list(shop_data.keys())}")

            # Process the document
            log("Starting document processing...")
            replace_placeholders_and_format(template_path, shop_data, output_path)
            log("Document processing completed successfully")
            
        except Exception as e:
            error_msg = f"Error processing document: {str(e)}"
            log(error_msg)
            log(f"Error type: {type(e).__name__}")
            if hasattr(e, 'args'):
                log(f"Error args: {e.args}")
            raise
            
    except Exception as e:
        error_msg = f"Fatal error: {str(e)}"
        log(error_msg)
        print(error_msg, file=sys.stderr)
        sys.exit(1) 